import requests, uuid, json
from docx import Document
from google.colab import userdata

subscription_key = userdata.get('KEY1')
endpoint = userdata.get('endpoint')
location = 'eastus2'

language_destination = 'pt-br'

def translator_text(text, target_language):
  path = '/translate'
  constructed_url = endpoint + path
  headers = {
      'Ocp-Apim-Subscription-Key': subscription_key, 
      'Ocp-Apim-Subscription-Region': location,
      'Content-type': 'application/json',
      'X-ClientTraceId': str(uuid.uuid4())
  }

  body = [{
      'text': text
  }]

  params = {
      'api-version': '3.0',
      'from' : 'en',
      'to': [target_language]
  }

  request = requests.post(constructed_url, params=params, headers=headers, json=body)
  response = request.json()
  return response[0]["translations"][0]["text"]

def translate_document(path):
  document = Document(path)
  full_text = []
  for paragraph in document.paragraphs:
    translated_text = translator_text(paragraph.text, language_destination)
    full_text.append(translated_text)

  translated_doc = Document()
  for line in full_text:    
    translated_doc.add_paragraph(line)
  path_translated = path.replace(".docx", f"_{language_destination}.docx")
  translated_doc.save(path_translated)
  return path_translated


# translator_text("O Captain! my Captain! our fearful trip is done", language_destination)
input_file = ""
# translate_document(input_file)
